import streamlit as st
import base64
import re
import io
from docx import Document
from PyPDF2 import PdfReader
from PIL import Image, ImageFilter
import pytesseract

# Page config
st.set_page_config(page_title="RE-DACT AI", layout="centered")

# Custom CSS
def local_css():
    st.markdown("""
        <style>
        html, body, [class*="css"]  {
            background: radial-gradient(circle at top, #001f3f, #000000);
            color: #ffffff;
            font-family: 'Segoe UI', sans-serif;
            text-align: center;
        }
        .title-text {
            font-size: 3.5em;
            font-weight: bold;
            margin-top: 80px;
        }
        .subtitle-text {
            font-size: 1.5em;
            color: #cccccc;
            margin-bottom: 40px;
        }
        .button-style {
            background-color: #f0ff00;
            color: black;
            border-radius: 999px;
            padding: 0.75em 2em;
            font-size: 1.1em;
            font-weight: bold;
            text-decoration: none;
        }
        .stButton>button {
            background-color: #f0ff00;
            color: black;
            font-weight: bold;
            border-radius: 999px;
            padding: 0.5em 1.5em;
            border: none;
        }
        </style>
    """, unsafe_allow_html=True)

local_css()

# UI header
st.markdown('<div class="title-text">The AI assistant that works everywhere</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-text">Available on web & desktop to protect your sensitive information</div>', unsafe_allow_html=True)

# Upload section
st.subheader("🔐 Upload your file")
uploaded_file = st.file_uploader("Choose a Word (.docx), PDF, or Image", type=["docx", "pdf", "png", "jpg", "jpeg"])

# Regex for masking
PATTERNS = {
    "Emails": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
    "Phone Numbers": r'\b\d{10}\b',
    "Account Numbers": r'\b\d{11,14}\b',
    "IFSC Codes": r'\b[A-Z]{4}\d{7}\b'
}

# Functions
def mask_text(paragraphs, selected):
    masked = []
    for para in paragraphs:
        for option in selected:
            para = re.sub(PATTERNS[option], '[DATA HIDDEN]', para)
        masked.append(para)
    return masked

def extract_pdf(file):
    reader = PdfReader(file)
    return [page.extract_text() for page in reader.pages]

def extract_docx(file):
    doc = Document(file)
    return [para.text for para in doc.paragraphs]

def extract_image_text(image):
    pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"  # adjust if needed
    return pytesseract.image_to_string(image)

def mask_image(image, selected):
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    for i, word in enumerate(data["text"]):
        for sel in selected:
            if re.match(PATTERNS[sel], word):
                x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
                region = image.crop((x, y, x + w, y + h))
                image.paste(region.filter(ImageFilter.GaussianBlur(5)), (x, y))
    return image

def to_docx(paragraphs):
    buffer = io.BytesIO()
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Options
if uploaded_file:
    st.divider()
    option = st.multiselect("What do you want to redact?", list(PATTERNS.keys()))
    if option:
        if uploaded_file.type == "application/pdf":
            original = extract_pdf(uploaded_file)
            masked = mask_text(original, option)
            st.subheader("📄 Masked Output:")
            st.write("\n".join(masked))
            docx = to_docx(masked)
            st.download_button("📥 Download Redacted DOCX", docx, file_name="masked_output.docx")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            original = extract_docx(uploaded_file)
            masked = mask_text(original, option)
            st.subheader("📄 Masked Output:")
            st.write("\n".join(masked))
            docx = to_docx(masked)
            st.download_button("📥 Download Redacted DOCX", docx, file_name="masked_output.docx")
        elif uploaded_file.type.startswith("image/"):
            image = Image.open(uploaded_file).convert("RGB")
            masked_img = mask_image(image, option)
            st.image(masked_img, caption="Image with sensitive data hidden", use_column_width=True)
            img_io = io.BytesIO()
            masked_img.save(img_io, format="PNG")
            st.download_button("📥 Download Blurred Image", img_io.getvalue(), file_name="blurred.png")
    else:
        st.info("Select at least one data type to redact.")
